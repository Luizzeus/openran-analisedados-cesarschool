# -*- coding: utf-8 -*-
"""
Gera o guia do palestrante em .docx a partir do briefing em linguagem leiga,
com a numeração dos slides de `apresentacao/seminario-g3-latencia-v2.pdf` (18 slides)
amarrada a cada trecho do texto para conduzir quem apresenta.

Saída: briefing-latencia-proxy-qoe.docx  (mesma pasta)
Uso:   python gerar_briefing_doc.py
Dependência: python-docx
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
OUT = HERE / "briefing-latencia-proxy-qoe.docx"

ACCENT = RGBColor(0x0E, 0x7C, 0x86)
MUTED = RGBColor(0x5A, 0x66, 0x72)

# ---------------------------------------------------------------------------
# Conteúdo: cada bloco = 1 slide da apresentação v2
#   (n, kicker, titulo do slide, "na tela", [parágrafos do roteiro], tempo aprox.)
# ---------------------------------------------------------------------------
SLIDES = [
    (1, "Capa", "Latência de rádio como proxy de QoE",
     "Título do trabalho e uma faixa com a linha do atraso das 100 medições.",
     ["Apresente a equipe e resuma o tema em uma frase: investigamos quando o "
      "\"tempo de espera\" dos dados na rede de celular fica alto o suficiente "
      "para sugerir que o usuário está tendo uma experiência ruim — e propusemos "
      "uma regra automática para reagir a isso, testada apenas em simulação."],
     "0:45"),

    (2, "O problema", "A pergunta do grupo",
     "A pergunta central e quatro números: 100 medições, 3 fases, 3 métricas, a palavra \"proxy\".",
     ["Explique latência com uma imagem do dia a dia: cada pedacinho de dado leva "
      "um tempinho para ir da antena até o aparelho; quando esse tempo sobe, vídeo "
      "trava, chamada fica robótica, jogo \"laga\".",
      "Explique o \"proxy\": no laboratório não existe uma nota de satisfação do "
      "usuário (o \"de 1 a 5, como foi sua ligação?\"). Então usamos o atraso como "
      "substituto — atraso alto costuma andar junto com usuário insatisfeito. "
      "Deixe claro: todo o trabalho fala desse substituto, não de satisfação medida."],
     "1:15"),

    (3, "Os dados", "Origem e forma dos dados",
     "Linha do tempo com as três fases: baseline (20), stress (60), recovery (20).",
     ["São 100 medições de um laboratório simulado — não é rede de operadora real "
      "e não há dados de pessoas.",
      "Elas cobrem três momentos: baseline = rede tranquila; stress = rede sob "
      "carga pesada, a \"hora do rush\"; recovery = logo depois do pico, a rede se "
      "recuperando."],
     "0:45"),

    (4, "Arquitetura", "Onde os dados ficam — bronze / silver / gold",
     "Diagrama de três caixas ligadas: dado cru → dado organizado → indicadores.",
     ["Slide mais técnico — passe rápido. A ideia: guardamos os dados em três "
      "camadas — uma cópia crua, uma versão organizada para consulta e a versão "
      "final já com os indicadores.",
      "Se perguntarem: não usamos bancos \"de tempo real\" porque os dados são um "
      "pacote fechado de 100 medições, não um fluxo ao vivo."],
     "0:40"),

    (5, "Qualidade", "Checagem antes de qualquer indicador",
     "Barras por fase separando instantes com tráfego dos instantes ociosos.",
     ["Antes de qualquer conta, conferimos os dados: sem furos, sem duplicatas.",
      "Achado importante: em mais da metade dos instantes de baseline e recovery "
      "não havia tráfego nenhum acontecendo. Nesses instantes o atraso aparece "
      "como zero — e zero aqui significa \"não havia nada para medir\", não "
      "\"estava perfeito\". Separamos esses instantes para não distorcer as contas."],
     "1:00"),

    (6, "Indicador 1", "KPI 1 — Atraso RLC por fase (mediana e p95)",
     "Linha do atraso nas 100 medições, colorida por fase, com a linha de corte tracejada.",
     ["Primeiro indicador: o atraso típico (a mediana, o \"meio\" das medições) e "
      "o atraso nos piores 5% dos instantes (o pico que se repete).",
      "Considerando só os instantes com tráfego real: no baseline o típico é "
      "~137; no stress ~159; no recovery ~126. Já os piores 5%: baseline ~204, "
      "stress ~191 e recovery ~438 — o recovery tem a pior cauda de todas."],
     "1:10"),

    (7, "Indicador 1", "KPI 1 — Distribuição do atraso por fase",
     "Uma \"caixa\" (boxplot) por fase mostrando onde os valores se concentram.",
     ["No stress a nuvem de pontos é compacta e alta — degradação consistente, "
      "sempre por perto do mesmo patamar.",
      "No baseline e no recovery a caixa fica \"esmagada\" perto do zero por causa "
      "dos instantes ociosos, mas com uma cauda que alcança o patamar do stress. "
      "O recovery tem dois picos altíssimos, acima do pior caso do stress."],
     "1:00"),

    (8, "Indicador 2", "KPI 2 — Fração de tempo em degradação",
     "A fórmula do segundo indicador em destaque e a definição da \"linha de corte\".",
     ["Segundo indicador: de todo o período, quanto tempo a rede passou acima de "
      "uma linha de corte que separa \"tranquila\" de \"congestionada\".",
      "A escolha inicial da linha é um valor estatístico do baseline (~105). "
      "Anuncie que os próximos três slides existem só para mostrar por que dá "
      "para confiar nessa linha."],
     "0:50"),

    (9, "Validação do limiar", "Por que 95 ≤ L < 133,7 µs dá o mesmo resultado",
     "Os valores observados espalhados numa régua, com uma faixa vazia destacada.",
     ["Entre os valores que realmente aconteceram existe um \"vão vazio\": nenhuma "
      "medição caiu entre ~95 e ~134. Nossa linha está dentro desse vão.",
      "Use a analogia do interruptor com folga: dá para mexer um pouco na posição "
      "da linha sem mudar o resultado. Qualquer linha nessa faixa leva à mesma "
      "conclusão."],
     "1:00"),

    (10, "Validação do limiar", "Sensibilidade — o critério muda a conclusão?",
     "Curvas mostrando o resultado para dezenas de posições diferentes da linha.",
     ["Testamos dezenas de linhas diferentes. Enquanto a linha fica na faixa de "
      "folga, a conclusão — \"só o stress está ruim\" — não muda.",
      "As linhas que rejeitamos são as que caem dentro do próprio stress: elas "
      "escondem a degradação, o que seria enganoso."],
     "0:50"),

    (11, "Validação do limiar", "O ponto de percentil é instável — a decisão não é",
     "Histograma do valor da linha após milhares de reembaralhamentos dos dados.",
     ["Reembaralhamos os dados milhares de vezes para ver o quanto o número da "
      "linha varia. Ele varia bastante — é pouca amostra.",
      "O ponto-chave da honestidade do trabalho: mesmo o número oscilando, a "
      "decisão que ele gera continua a mesma. Admitimos a fragilidade do número e "
      "mostramos que ela não afeta a resposta."],
     "0:50"),

    (12, "Resultado", "KPI 2 — resultados e fórmula final",
     "Barras do tempo acima da linha por fase e a fórmula final dos dois indicadores.",
     ["Resultado do segundo indicador: baseline 25% do tempo acima da linha, "
      "stress 100%, recovery 25%. No pico de carga, todas as medições estavam "
      "acima da linha.",
      "E a regra de persistência: não basta um tropeço isolado — tem que ficar "
      "acima da linha por 5 medições seguidas. Isso acontece 56 vezes no stress e "
      "nenhuma vez no baseline ou no recovery."],
     "1:00"),

    (13, "Cruzamento", "Atraso e vazão — contexto operacional do proxy",
     "Gráfico de dispersão cruzando o atraso com a quantidade de dados trafegando.",
     ["Cruzamos o atraso com o volume de dados em trânsito. O atraso alto "
      "acompanha as janelas de maior atividade da rede.",
      "Faça a ressalva: associação não é causa, e isso não confirma experiência "
      "ruim — são medidas de coisas diferentes. É contexto, não prova."],
     "0:50"),

    (14, "Recomendação", "Política A1 candidata — execução simulada (dry-run)",
     "Fluxograma: fase → duas condições → decisão \"agir\" ou \"observar\".",
     ["Nossa recomendação: durante o pico de carga, acionar priorização de "
      "tráfego e investigação da sessão — na prática, dar preferência a quem está "
      "sofrendo e olhar o que está acontecendo.",
      "O gatilho exige duas condições ao mesmo tempo: a rede ruim na maior parte "
      "do tempo E isso sustentado por 5 medições seguidas. Só o stress satisfaz "
      "as duas.",
      "Frase obrigatória: tudo isso rodou em modo simulação. Nenhuma configuração "
      "de rede real foi alterada; nada foi enviado a equipamento de verdade."],
     "1:15"),

    (15, "Modelo comparativo", "EWMA — controle estatístico como gatilho de persistência",
     "Uma \"carta de controle\": a linha do EWMA subindo e cruzando um limite de alerta.",
     ["Para não depender de um único método, comparamos com o EWMA — pense num "
      "\"termômetro de memória curta\", que dá mais peso ao que acabou de "
      "acontecer.",
      "Ele confirma o stress como o momento crítico e, de bônus, acende alerta em "
      "55% do recovery — ou seja, aponta que a rede ainda não voltou totalmente "
      "ao normal depois do pico."],
     "1:00"),

    (16, "Comparativo", "Três modelos, a mesma decisão de fase — por caminhos diferentes",
     "Barras comparando, por fase, o que cada um dos três modelos sinaliza.",
     ["Três métodos independentes concordam no essencial: agir só no stress. Essa "
      "convergência é o que dá confiança na conclusão.",
      "Cada um chega lá por um caminho: o modelo pronto do laboratório detecta o "
      "\"evento de carga\" (vários sinais subindo juntos), não a latência em si, e "
      "é cego à cauda do recovery; o modelo do grupo mede latência numa unidade "
      "com significado físico; o EWMA serve para vigiar a recuperação."],
     "1:00"),

    (17, "Limitações", "O que estes números não dizem",
     "Lista das limitações do estudo.",
     ["Seja direto com o que o trabalho NÃO diz: não é rede real (é laboratório "
      "simulado); não é satisfação medida (atraso é substituto, ninguém "
      "perguntou ao usuário); poucos aparelhos e poucas medições; o recovery não "
      "voltou 100% ao normal; e a recomendação é um ensaio, não algo para ligar "
      "numa rede em produção."],
     "0:50"),

    (18, "Conclusão", "O que os dados sustentam",
     "Os pontos que os dados sustentam e a mensagem final.",
     ["Feche com a conclusão em uma linha: nos dados analisados, o atraso só "
      "indica experiência potencialmente ruim quando deixa de ser um pico isolado "
      "e passa a dominar o período — e isso, aqui, acontece exclusivamente "
      "durante o pico de carga.",
      "Lembre que tudo é reproduzível: rodar o notebook do zero chega aos mesmos "
      "números e gráficos."],
     "0:40"),
]

# ---------------------------------------------------------------------------
doc = Document()

base = doc.styles["Normal"]
base.font.name = "Calibri"
base.font.size = Pt(11)

def _spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)

# ---- título ----
t = doc.add_paragraph()
r = t.add_run("Guia do palestrante")
r.bold = True
r.font.size = Pt(20)
t.paragraph_format.space_after = Pt(2)

s = doc.add_paragraph()
r = s.add_run("Latência de rádio como proxy de QoE  ·  Projeto Integrador G3  ·  Módulo 09 — CESAR School")
r.font.size = Pt(11)
r.font.color.rgb = MUTED

e = doc.add_paragraph()
r = e.add_run("Equipe: Carlos Alberto · Éverton Gomes · Gerson Francisco · Luiz Carlos Santos")
r.font.size = Pt(10)
r.font.color.rgb = MUTED

doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run(
    "Este guia acompanha, slide a slide, a apresentação "
).font.size = Pt(11)
intro.add_run("apresentacao/seminario-g3-latencia-v2.pdf").italic = True
intro.add_run(
    " (18 slides). Para cada slide há o que aparece na tela, um roteiro em "
    "linguagem simples (pode ser parafraseado, não precisa ser lido) e um tempo "
    "aproximado. Total sugerido: cerca de 15 a 17 minutos — ajuste conforme o "
    "tempo disponível."
).font.size = Pt(11)

# ---- roteiro-resumo (tabela) ----
doc.add_paragraph()
h = doc.add_paragraph()
r = h.add_run("Visão geral")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = ACCENT

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
hdr = tbl.rows[0].cells
for i, txt in enumerate(("Slide", "Bloco", "Tema", "≈ tempo")):
    hdr[i].text = txt
    for pr in hdr[i].paragraphs:
        for rr in pr.runs:
            rr.bold = True
for n, kicker, titulo, _tela, _rot, tempo in SLIDES:
    row = tbl.add_row().cells
    row[0].text = str(n)
    row[1].text = kicker
    row[2].text = titulo
    row[3].text = tempo

# ---- corpo, slide a slide ----
doc.add_page_break()
body = doc.add_paragraph()
r = body.add_run("Roteiro slide a slide")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = ACCENT
doc.add_paragraph()

for n, kicker, titulo, tela, roteiro, tempo in SLIDES:
    hp = doc.add_paragraph()
    hp.paragraph_format.space_before = Pt(10)
    hp.paragraph_format.space_after = Pt(2)
    rr = hp.add_run(f"Slide {n:02d}")
    rr.bold = True
    rr.font.size = Pt(14)
    rr.font.color.rgb = ACCENT
    rr2 = hp.add_run(f"   {kicker} — {titulo}")
    rr2.bold = True
    rr2.font.size = Pt(12)
    rr3 = hp.add_run(f"   ({tempo})")
    rr3.font.size = Pt(10)
    rr3.font.color.rgb = MUTED

    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(2)
    rt = tp.add_run("Na tela: ")
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.color.rgb = MUTED
    rv = tp.add_run(tela)
    rv.font.size = Pt(10)
    rv.font.color.rgb = MUTED

    lp = doc.add_paragraph()
    lp.paragraph_format.space_after = Pt(1)
    rl = lp.add_run("Roteiro:")
    rl.bold = True
    rl.font.size = Pt(10.5)

    for par in roteiro:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(3)
        bp.add_run(par).font.size = Pt(11)

doc.add_paragraph()
foot = doc.add_paragraph()
rf = foot.add_run(
    "Texto de apoio (versão contínua, sem marcação de slides): "
)
rf.font.size = Pt(9.5)
rf.font.color.rgb = MUTED
rf2 = foot.add_run("briefing-latencia-proxy-qoe.md")
rf2.italic = True
rf2.font.size = Pt(9.5)
rf2.font.color.rgb = MUTED

doc.save(str(OUT))
print("DOCX:", OUT, f"({OUT.stat().st_size/1024:.0f} KB)")
